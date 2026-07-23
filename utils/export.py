from io import BytesIO
from datetime import datetime

from docx import Document
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate


def export_markdown(chat_history):
    """
    Returns Markdown text.
    """

    lines = [
        "# AI PDF Research Assistant",
        "",
        f"Generated: {datetime.now().strftime('%d %B %Y %H:%M:%S')}",
        "",
        "---",
        ""
    ]

    for question, answer in chat_history:

        lines.append("## 👤 User")
        lines.append(question)
        lines.append("")

        lines.append("## 🤖 Assistant")
        lines.append(answer)
        lines.append("")

    return "\n".join(lines)


def export_docx(chat_history):
    """
    Returns DOCX as bytes.
    """

    document = Document()

    document.add_heading(
        "AI PDF Research Assistant",
        level=1
    )

    document.add_paragraph(
        f"Generated: {datetime.now().strftime('%d %B %Y %H:%M:%S')}"
    )

    for question, answer in chat_history:

        document.add_heading(
            "User",
            level=2
        )

        document.add_paragraph(question)

        document.add_heading(
            "Assistant",
            level=2
        )

        document.add_paragraph(answer)

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return buffer.getvalue()


def export_pdf(chat_history):
    """
    Returns PDF as bytes.
    """

    buffer = BytesIO()

    pdf = SimpleDocTemplate(buffer)

    styles = getSampleStyleSheet()

    elements = []

    elements.append(
        Paragraph(
            "<b>AI PDF Research Assistant</b>",
            styles["Title"]
        )
    )

    elements.append(
        Paragraph(
            datetime.now().strftime(
                "Generated: %d %B %Y %H:%M:%S"
            ),
            styles["Normal"]
        )
    )

    for question, answer in chat_history:

        elements.append(
            Paragraph(
                "<b>User</b>",
                styles["Heading2"]
            )
        )

        elements.append(
            Paragraph(
                question,
                styles["BodyText"]
            )
        )

        elements.append(
            Paragraph(
                "<b>Assistant</b>",
                styles["Heading2"]
            )
        )

        elements.append(
            Paragraph(
                answer,
                styles["BodyText"]
            )
        )

    pdf.build(elements)

    buffer.seek(0)

    return buffer.getvalue()